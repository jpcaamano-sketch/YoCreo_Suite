"""
Funciones de exportacion para YoCreo Suite
Protocolo Estandar v2.0
- PDFs con reportlab en memoria
- Boton copiar HTML/JS ancho completo
- Encabezado con boton de manual
"""

import io
import os
import base64
import streamlit as st
import streamlit.components.v1 as components
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from docx import Document

# Ruta de manuales
MANUALES_PATH = os.path.join(os.path.dirname(__file__), "..", "manuales")


# ==================== ENCABEZADO CON MANUAL ====================

def render_encabezado(practica_key, titulo, descripcion):
    """
    Renderiza el encabezado de una practica con boton de manual

    Args:
        practica_key: Key de la practica (ej: 'priorizador_tareas')
        titulo: Titulo de la practica
        descripcion: Descripcion breve
    """
    # Verificar si existe el manual
    pdf_path = os.path.join(MANUALES_PATH, f"{practica_key}.pdf")
    tiene_manual = os.path.exists(pdf_path)

    if tiene_manual:
        col1, col2 = st.columns([0.85, 0.15])
        with col1:
            st.markdown(f"### {titulo}")
        with col2:
            if st.button("📖", key=f"manual_btn_{practica_key}", help=f"Ver Manual de {titulo}"):
                # Abrir PDF en nueva pestaña
                with open(pdf_path, "rb") as f:
                    pdf_base64 = base64.b64encode(f.read()).decode()

                js_code = f'''
                <script>
                    const base64Data = "{pdf_base64}";
                    const byteCharacters = atob(base64Data);
                    const byteNumbers = new Array(byteCharacters.length);
                    for (let i = 0; i < byteCharacters.length; i++) {{
                        byteNumbers[i] = byteCharacters.charCodeAt(i);
                    }}
                    const byteArray = new Uint8Array(byteNumbers);
                    const blob = new Blob([byteArray], {{type: 'application/pdf'}});
                    const blobUrl = URL.createObjectURL(blob);
                    window.open(blobUrl, '_blank');
                </script>
                '''
                components.html(js_code, height=0, width=0)
    else:
        st.markdown(f"### {titulo}")

    st.write(descripcion)


# ==================== BOTON COPIAR (Protocolo Estandar) ====================

def copy_button_component(text_to_copy, key="copy"):
    """
    Muestra boton copiar HTML/JS ancho completo (Naranja)
    Segun protocolo: debajo del area de texto
    """
    if not text_to_copy:
        return

    text_js = text_to_copy.replace('`', '\\`').replace('${', '\\${').replace('\n', '\\n').replace('\r', '')

    component_html = f"""
    <head>
        <link href="https://fonts.googleapis.com/css2?family=Inter:wght@400;600;700&display=swap" rel="stylesheet">
    </head>
    <style>
        body {{ margin: 0; padding: 0; }}
        .copy-btn {{
            width: 100%;
            background-color: #FF6B4E;
            color: white;
            border: none;
            border-radius: 8px;
            padding: 0.75rem;
            font-family: 'Inter', sans-serif;
            font-weight: 700;
            font-size: 16px;
            cursor: pointer;
            box-shadow: 0 4px 6px rgba(0,0,0,0.1);
            transition: transform 0.1s;
        }}
        .copy-btn:hover {{
            transform: translateY(-2px);
            background-color: #e65539;
        }}
    </style>
    <button class="copy-btn" id="btn_{key}" onclick="copyToClipboard_{key}()">Copiar al Portapapeles</button>
    <script>
        function copyToClipboard_{key}() {{
            const text = `{text_js}`;
            navigator.clipboard.writeText(text).then(() => {{
                const btn = document.getElementById('btn_{key}');
                btn.innerText = 'Copiado';
                btn.style.backgroundColor = '#28a745';
                setTimeout(() => {{
                    btn.innerText = 'Copiar al Portapapeles';
                    btn.style.backgroundColor = '#FF6B4E';
                }}, 2000);
            }});
        }}
    </script>
    """
    components.html(component_html, height=60)


def text_area_with_copy(label, text, key, height=300):
    """
    Muestra un text_area editable grande (editor unico)
    El boton copiar se agrega por separado con copy_button_component

    Args:
        label: Etiqueta del campo
        text: Texto inicial
        key: Key unico para el componente
        height: Altura del text_area

    Returns:
        str: El texto editado
    """
    edited_text = st.text_area(label, value=text, height=height, key=key)
    return edited_text


# ==================== PDF con ReportLab ====================

def create_pdf_reportlab(titulo, secciones):
    """
    Crea un documento PDF profesional con reportlab en memoria

    Args:
        titulo: Titulo del documento
        secciones: Lista de tuplas (titulo_seccion, contenido)

    Returns:
        bytes: Contenido del archivo PDF
    """
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    elements = []

    styles = getSampleStyleSheet()

    # Estilo titulo
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        alignment=TA_CENTER,
        fontSize=18,
        spaceAfter=20
    )

    # Estilo subtitulo seccion
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=14,
        textColor='#4E32AD',
        spaceAfter=10
    )

    # Estilo normal
    normal_style = ParagraphStyle(
        'CustomNormal',
        parent=styles['Normal'],
        fontSize=11,
        leading=14,
        alignment=TA_LEFT
    )

    # Titulo principal
    elements.append(Paragraph(titulo, title_style))
    elements.append(Spacer(1, 20))

    # Secciones
    for seccion_titulo, contenido in secciones:
        elements.append(Paragraph(seccion_titulo, heading_style))
        texto_formateado = contenido.replace("\n", "<br/>")
        elements.append(Paragraph(texto_formateado, normal_style))
        elements.append(Spacer(1, 15))

    doc.build(elements)
    buffer.seek(0)
    return buffer.getvalue()


# ==================== WORD ====================

def create_word_document(titulo, secciones):
    """
    Crea un documento Word con secciones

    Args:
        titulo: Titulo del documento
        secciones: Lista de tuplas (titulo_seccion, contenido)

    Returns:
        bytes: Contenido del archivo Word
    """
    doc = Document()
    doc.add_heading(titulo, 0)

    for seccion_titulo, contenido in secciones:
        doc.add_heading(seccion_titulo, level=1)
        doc.add_paragraph(contenido)

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


# ==================== UTILIDADES ====================

def get_pdf_mime():
    """Retorna el MIME type para PDF"""
    return "application/pdf"


def get_word_mime():
    """Retorna el MIME type para Word"""
    return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def show_download_section(st, contenido_texto, nombre_default="documento"):
    """
    Muestra la seccion de descarga estandar (CAJA 4)
    Sin emojis, con selector de formato

    Args:
        st: Modulo streamlit
        contenido_texto: Texto completo a exportar
        nombre_default: Nombre por defecto del archivo
    """
    col1, col2 = st.columns(2)
    with col1:
        nombre_archivo = st.text_input(
            "Nombre del archivo",
            value=nombre_default
        )
    with col2:
        formato = st.selectbox(
            "Formato",
            ["PDF", "Texto (.txt)"]
        )

    if formato == "PDF":
        pdf_data = create_pdf_reportlab(
            nombre_default.replace("_", " ").title(),
            [("Contenido", contenido_texto)]
        )
        st.download_button(
            "Descargar PDF",
            data=pdf_data,
            file_name=f"{nombre_archivo}.pdf",
            mime="application/pdf",
            use_container_width=True
        )
    else:
        st.download_button(
            "Descargar TXT",
            data=contenido_texto,
            file_name=f"{nombre_archivo}.txt",
            mime="text/plain",
            use_container_width=True
        )
